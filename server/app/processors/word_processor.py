# app/processors/word_processor.py
import io
import time
import logging
from docx import Document
from app.processors.base_processor import BaseProcessor

logger = logging.getLogger(__name__)

class WordProcessor(BaseProcessor):
    """Word document processor"""
    
    def get_supported_extensions(self) -> set:
        return {'docx', 'doc'}
    
    
    def supports_format(self, file_format: str) -> bool:
        """Check if the processor supports the given file format"""
        return file_format.lower() in self.get_supported_extensions()
    
    
    async def process(self, file_content: bytes, filename: str, **kwargs) -> str:
        """Parse Word documents"""
        logger.info("Starting Word document parsing")
        start_time = time.time()
        
        try:
            document = Document(io.BytesIO(file_content))
            text_parts = []
            
            # Extract paragraphs
            for para in document.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text.strip())
            
            # Extract tables with proper structure
            for table in document.tables:
                table_text = ["TABLE DATA:"]
                for row in table.rows:
                    row_data = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_data.append(cell.text.strip())
                    if row_data:
                        table_text.append(" | ".join(row_data))
                if len(table_text) > 1:
                    text_parts.append("\n".join(table_text))
            
            final_text = "\n\n".join(text_parts)
            logger.info(f"Word document parsing completed in {time.time() - start_time:.2f}s. Text length: {len(final_text)}")
            return final_text
            
        except Exception as e:
            logger.error(f"Word document parsing failed: {str(e)}")
            raise Exception(f"Word document parsing failed: {str(e)}")